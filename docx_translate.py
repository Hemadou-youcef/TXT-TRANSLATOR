import sys
import os
import time
import docx
from docx import Document
from docx.shared import Pt
from docx.shared import Length
from docx.shared import Inches, Cm

from urllib.parse import urlparse
import txt_translate
from progress.bar import IncrementalBar
import shutil
import re


class Docx_trans:
    def Docx(self):
        try:
            while True:
                try:
                    # self.novel_name = input("WRITE FILE NAME: ").upper()
                    # if not self.novel_name :
                    #     self.novel_name="NOVEL"
                    # for key in ['\\','/',':','?','*','<','>','|','"']:
                    #     self.novel_name = self.novel_name.replace(key,'')

                    self.novel_name = ""
                    self.to = input("TO?:(DEFAULT=EN): ").lower()
                    if not self.to :
                        self.to="EN"   
                    Acc = txt_translate.txt_translate()
                    

                    os.chdir("txt_files/" + self.novel_name)
                    
                    self.complete_name=["".join(f.split('.')[:-1]) for f in os.listdir() if f.endswith(".docx")]
                    self.complete_name.sort(key=self.natural_keys)
                    
                    
                    for novel in self.complete_name:
                        os.system('cls')
                        print(novel.upper())
                        os.mkdir(novel)
                        shutil.move(os.getcwd() + "/" + novel + ".docx",os.getcwd() + "/" + novel)
                        os.chdir(novel)
                        self.cutter()
                        os.chdir("../..")
                        Acc.initialisation(argv=[novel,"auto",self.to,"txt",10])
                        os.chdir("..")
                        self.merge(novel)
                        os.chdir("txt_files/" + self.novel_name)
                    break
                except Exception as e:
                    print(e)

            print("DONE!!")
            print("THANK FOR USING OUR PRODUCT")
        except Exception as e:
            print("ERROR ENTRING LIGHT NOVEL")
            print(e)
    def cutter(self):
        try:
            self.temp_txt = ""
            self.range = 2500

            file_name=[f for f in os.listdir() if f.endswith(".docx")]
            file_name.sort(key=self.natural_keys)

            for path in file_name:
                text = self.getText(path)
                
                txt = text.split("\n")
                txt_read_lenght = len(text)

                i = 1
                if txt_read_lenght > self.range:
                    for line in txt:
                        self.temp_txt += line + "\n"
                        if len(self.temp_txt) > self.range:
                            all = open(str(i) + ".txt", "w",encoding="utf-8")
                            all.write(self.temp_txt)
                            all.close()
                            self.temp_txt = ""
                            i += 1
                    if len(self.temp_txt) >= 1:
                        all = open(str(i) + ".txt", "w",encoding="utf-8")
                        all.write(self.temp_txt)
                        all.close()
                else:
                    all = open(str(i) + ".txt", "w",encoding="utf-8")
                    all.write(text)
                    all.close()
                os.remove(path) 
            
        except Exception as e:
            print("ERROR ENTRING LIGHT NOVEL")
            print(e)
    def merge(self,name):
        try:
            os.chdir("translated_novel/" + name)
            self.txt = ""

            file_name=[f for f in os.listdir() if f.endswith(".txt")]
            file_name.sort(key=self.natural_keys)

            for path in file_name:
                txt = open(path, "r")
                self.txt +=  txt.read() + "\n"
                txt.close()
                os.remove(path) 

            all = open(name + ".txt", "w")
            all.write(self.txt)
            all.close()

            document = Document()
            sections = document.sections
            for section in sections:
                section.top_margin = Cm(2.54)
                section.bottom_margin = Cm(2.54)
                section.left_margin = Cm(2.65)
                section.right_margin = Cm(2.65)
            style = document.styles['Normal']
            font = style.font
            font.name = 'Courier New'
            font.size = Pt(10.5)
            if self.to.lower() == "ar":
                font.rtl = True
            style.paragraph_format.line_spacing = 1.15
            myfile = open(name + ".txt",encoding="utf-8").read()

            if self.to.lower() != "ar":
                myfile = re.sub(r'[^\x00-\x7F]+|\x0c',' ', myfile) # remove all non-XML-compatible characters
            p = document.add_paragraph(myfile)
            document.save(name + '.docx')
            os.remove(name + ".txt") 
            os.chdir("../..")
            
        except Exception as e:
            print("ERROR ENTRING LIGHT NOVEL")
            print(e)
    def atoi(self,text):
        return int(text) if text.isdigit() else text
    def natural_keys(self,text):
        return [ self.atoi(c) for c in re.split(r'(\d+)', text) ]
    def getText(self,filename):
        doc = docx.Document(filename)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        return '\n'.join(fullText)
os.system('cls')
Acc = Docx_trans()
Acc.Docx()
time.sleep(5)
sys.exit()